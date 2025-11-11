from docx import Document

import json
import os
import logging
import shlex
import subprocess
import platform
import tempfile

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
    handlers=[logging.StreamHandler()],
)
logger = logging.getLogger(__name__)

config = json.load(open("config.json", "r", encoding="utf-8"))

project_path = os.path.abspath(os.path.dirname(__file__))


def _get_drive_letter(path):
    """
    获取路径所在的卷（驱动器字母）
    
    Args:
        path: 文件或目录路径
        
    Returns:
        在Windows上返回驱动器字母（如 'C:'），在Linux/Mac上返回空字符串
    """
    if platform.system() == "Windows":
        drive, _ = os.path.splitdrive(os.path.abspath(path))
        return drive
    return ""


def _revert_docx_to_md(doc: Document, md_project_name: str) -> str:
    image_map = {}
    image_count = 0

    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            image_count += 1
            image_data = rel.target_part.blob
            ext = rel.target_part.partname.split("/")[-1].split(".")[-1].lower()
            if ext not in ["png", "jpg", "jpeg", "gif", "bmp", "webp"]:
                ext = "jpg"

            image_name = f"img_{image_count}.{ext}"
            image_path = os.path.join(config["MD_DIR_PATH"], md_project_name, "img", image_name)
            with open(image_path, "wb") as f:
                f.write(image_data)
            image_map[rel.rId] = f"./img/{image_name}"

    image_count = 0
    full_text = ""

    parent = doc.element.body if hasattr(doc, "element") else doc
    markdown_struct = []
    for child in parent.getchildren():
        if child.tag.endswith("}p"):
            markdown_struct.append("p")
        elif child.tag.endswith("}tbl"):
            markdown_struct.append("tb")

    tables = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", "<br>") for cell in row.cells]
            rows.append("| " + " | ".join(cells) + " |")

        if len(rows) > 1:
            sep = "|" + "|".join([" --- "] * len(table.columns)) + "|"
            rows.insert(1, sep)
        tables.append("\n".join(rows))

    paras = []
    for para in doc.paragraphs:
        para_content = ""
        for run in para.runs:
            for child in run._element.getchildren():
                if child.tag.endswith("}drawing") or child.tag.endswith("}pict"):
                    blips = child.xpath(".//*[local-name()='blip']")

                    if blips:
                        rId = blips[0].get(
                            "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                        )
                        if rId:
                            if rId in [r.rId for r in doc.part.rels.values()]:
                                image_count += 1
                                image_path = image_map.get(rId, f"img_{image_count}.jpg")
                                para_content += f"\n\n![图片{image_count}]({image_path})\n\n"

                elif child.tag.endswith("}t"):
                    text = child.text or ""
                    para_content += text

        paras.append(para_content)

    for item in markdown_struct:
        if item == "p":
            full_text += paras.pop(0) + "\n\n"
        elif item == "tb":
            full_text += tables.pop(0) + "\n\n"

    return full_text.strip()


def convert_docx_to_markdown(docx_path: str) -> str:
    if not os.path.exists(docx_path):
        return "文件不存在"

    if not docx_path.endswith(".docx"):
        return "文件不是docx文件"

    docx_name = os.path.basename(docx_path).split(".")[0]
    doc = Document(docx_path)
    todo_docx_dir = os.path.join(".", "todo", config['PROJECT_NAME'], docx_name)
    os.makedirs(todo_docx_dir, exist_ok=True)
    os.makedirs(os.path.join(todo_docx_dir, "img"), exist_ok=True)
    todo_md_path = os.path.join(todo_docx_dir, "todo.md")
    with open(todo_md_path, "w+", encoding="utf-8") as f:
        f.write(_revert_docx_to_md(doc, todo_docx_dir))


def convert_pdf_to_markdown(pdf_path: str) -> str:
    from PyPDF2 import PdfReader

    import os

    if not os.path.exists(pdf_path) or not os.path.isfile(pdf_path):
        return "pdf文件不存在"
    elif not pdf_path.endswith(".pdf"):
        return "文件不是pdf文件"
    else:
        with open(pdf_path, "rb") as f:
            reader = PdfReader(f)
            content = ""
            skipped_content_len = 0

            # 读取所有页面内容
            for page_num in range(len(reader.pages)):
                page = reader.pages[page_num]
                page_content = page.extract_text()
                content += page_content + "\n\n"

        pdf_name = os.path.basename(pdf_path).split(".")[0]
        todo_pdf_dir = os.path.join(".", "todo", config['PROJECT_NAME'], pdf_name)
        if not os.path.exists(todo_pdf_dir):
            os.makedirs(todo_pdf_dir, exist_ok=True)

        todo_md_path = os.path.join(todo_pdf_dir, "todo.md")
        with open(todo_md_path, "w+", encoding="utf-8") as f:
            f.write(content)

    return "pdf文件转换为markdown文件成功"


def _execute_script_subprocess(script_command, env_vars=None) -> str:
    """
    使用 subprocess 模块执行脚本（推荐）

    Args:
        script_command: 要执行的命令字符串
        env_vars: 要传递的环境变量字典，例如 {"CURSOR_API_KEY": "..."}
    """
    try:
        # 检测操作系统
        is_windows = platform.system() == "Windows"
        
        # 如果需要在命令前设置环境变量，可以在命令中导出
        todo_dir = os.path.join(project_path, "todo", config['PROJECT_NAME'])
        
        if is_windows:
            # Windows 使用 cmd /c
            # 获取目标目录所在的卷
            drive = _get_drive_letter(todo_dir)

            if drive:
                # 如果目标目录在不同卷，需要先切换到该卷
                base_command = rf"{drive} && cd {todo_dir}"
            else:
                base_command = rf"cd {todo_dir}"

            full_command = ""
            if env_vars:
                env_exports = " && ".join(
                    [f"set {k}={shlex.quote(str(v))}" for k, v in env_vars.items()]
                )
                full_command = f"{base_command} && {env_exports} && {script_command}"
            else:
                full_command = f"{base_command} && {script_command}"

            result = subprocess.run(
                ["cmd", "/c", full_command],
                capture_output=True,
                text=True,
                encoding="utf-8",
                errors="replace",
                check=True,
            )
        else:
            # Linux/Unix 使用 bash -c
            base_command = f"cd {shlex.quote(todo_dir.replace(os.sep, '/'))}"
            full_command = ""
            if env_vars:
                env_exports = " ".join(
                    [f"export {k}={shlex.quote(str(v))}" for k, v in env_vars.items()]
                )
                full_command = f"{base_command} && {env_exports} && {script_command}"
            else:
                full_command = f"{base_command} && {script_command}"
            
            result = subprocess.run(
                ["bash", "-c", full_command],
                capture_output=True,
                text=True,
                encoding="utf-8",
                errors="replace",
                check=True,
            )
        logger.info("执行成功！")
        logger.info(f"输出: {result.stdout}")
        return result.stdout
    except subprocess.CalledProcessError as e:
        logger.error("执行失败！")
        logger.error(f"错误: {e.stderr}")
        return "执行失败！"
    except Exception as e:
        logger.error(f"详细信息: {e}")
        return "执行失败！"


def analyze_what_to_do():
    env_vars = {"CURSOR_API_KEY": config["CURSOR_API_KEY"]}

    prompt = """
    不要等待任何提示！直接开始分析！
    请综合分析目录下所有文件内容，
    以需求点的形式一一列举，
    以markdown格式写到本文件夹下的todo.md文件中，写完需求即可不用执行更多操作！
    """

    opinion = ""
    opinion_file = os.path.join(".", "opinion", f"{config['PROJECT_NAME']}.md")
    if os.path.exists(opinion_file):
        with open(opinion_file, "r", encoding="utf-8") as f:
            opinion = f.read()

    if len(opinion) > 0:
        prompt += f"""

        审核员意见如下：
        {opinion}
        """

    development_log = ""
    development_log_file = os.path.join(".", "dist", config['PROJECT_NAME'], "development_log.md")
    if os.path.exists(development_log_file):
        with open(development_log_file, "r", encoding="utf-8") as f:
            development_log = f.read()

    if len(development_log) > 0:
        prompt += f"""

        分析中你必须考虑开发日志，并根据开发日志调整分析结果。
        开发日志如下：
        {development_log}
        """

    prompt += "\n\n注意！1. 分析中你必须并根据审核员意见和开发日志调整分析结果。\n2. 你不允许对所在目录的父目录进行写入操作！\n"

    if config["MOCK"]:
        execute_result = _execute_script_subprocess(
            f'python {config["SIM_CURSOR_PATH"]} -p --force --output-format text "{prompt}"',
            env_vars=env_vars
        )
    elif platform.system() == "Windows" and "EXECUTE_PATH" in config:
        with tempfile.NamedTemporaryFile(
            mode='w', 
            encoding='utf-8',
            delete=False, 
            suffix=".prompt",
            dir="."
        ) as temp_file:
            temp_file.write(prompt)
            temp_file_path = os.path.abspath(os.path.join(".", temp_file.name))
        execute_result = _execute_script_subprocess(
            f'{config["EXECUTE_PATH"]} -p --force --output-format text --prompt-file {temp_file_path}', env_vars=env_vars
        )
    else:
        execute_result = _execute_script_subprocess(
            f'{config["CURSOR_PATH"]} -p --force --output-format text "{prompt}"', env_vars=env_vars
        )

    return execute_result
